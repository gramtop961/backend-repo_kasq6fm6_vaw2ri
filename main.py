import os
import io
import uuid
from datetime import datetime, timedelta, timezone
from typing import Optional, List

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel, EmailStr
from passlib.context import CryptContext
import jwt

from database import db, create_document, get_documents
from schemas import User as UserSchema, Document as DocumentSchema, Summary as SummarySchema

from pymongo import ReturnDocument
from bson import ObjectId

from pypdf import PdfReader
from docx import Document as DocxDocument

import requests

JWT_SECRET = os.getenv("JWT_SECRET", "devsecret-change-me")
JWT_ALG = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login")

app = FastAPI(title="Gamma API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ------------------------
# Utils
# ------------------------

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALG)


def get_user_by_email(email: str):
    if db is None:
        return None
    return db["user"].find_one({"email": email})


def get_current_user(token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(status_code=401, detail="Invalid authentication")
        user = get_user_by_email(email)
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


# ------------------------
# Auth
# ------------------------
class SignupBody(BaseModel):
    name: str
    email: EmailStr
    password: str


@app.post("/auth/signup")
def signup(body: SignupBody):
    existing = get_user_by_email(body.email)
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    hashed = pwd_context.hash(body.password)
    data = UserSchema(name=body.name, email=body.email, password_hash=hashed)
    user_id = create_document("user", data)
    token = create_access_token({"sub": body.email})
    return {"access_token": token, "token_type": "bearer", "user_id": user_id}


@app.post("/auth/login")
def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = get_user_by_email(form_data.username)
    if not user:
        raise HTTPException(status_code=400, detail="Incorrect email or password")
    if not pwd_context.verify(form_data.password, user.get("password_hash")):
        raise HTTPException(status_code=400, detail="Incorrect email or password")
    token = create_access_token({"sub": user["email"]})
    return {"access_token": token, "token_type": "bearer"}


# ------------------------
# Health & schema
# ------------------------
@app.get("/")
def read_root():
    return {"service": "Gamma API", "status": "ok"}


@app.get("/schema")
def read_schema():
    return {
        "collections": ["user", "document", "summary"],
    }


@app.get("/test")
def test_database():
    response = {
        "backend": "✅ Running",
        "database": "❌ Not Available",
        "database_url": None,
        "database_name": None,
        "connection_status": "Not Connected",
        "collections": []
    }

    try:
        if db is not None:
            response["database"] = "✅ Available"
            response["database_url"] = "✅ Configured"
            response["database_name"] = db.name if hasattr(db, 'name') else "✅ Connected"
            response["connection_status"] = "Connected"
            try:
                collections = db.list_collection_names()
                response["collections"] = collections[:10]
                response["database"] = "✅ Connected & Working"
            except Exception as e:
                response["database"] = f"⚠️  Connected but Error: {str(e)[:50]}"
        else:
            response["database"] = "⚠️  Available but not initialized"
    except Exception as e:
        response["database"] = f"❌ Error: {str(e)[:50]}"

    response["database_url"] = "✅ Set" if os.getenv("DATABASE_URL") else "❌ Not Set"
    response["database_name"] = "✅ Set" if os.getenv("DATABASE_NAME") else "❌ Not Set"
    return response


# ------------------------
# Document upload & extraction
# ------------------------
class UploadResponse(BaseModel):
    document_id: str
    extracted_chars: int


def extract_text_from_file(file: UploadFile, bytes_data: bytes) -> str:
    if file.content_type == "application/pdf":
        reader = PdfReader(io.BytesIO(bytes_data))
        text_parts: List[str] = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts).strip()
    elif file.content_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
        doc = DocxDocument(io.BytesIO(bytes_data))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        try:
            return bytes_data.decode("utf-8")
        except Exception:
            return ""


@app.post("/documents/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    visibility: str = Form("private"),
    token: str = Depends(oauth2_scheme)
):
    user = get_current_user(token)
    raw = await file.read()
    text = extract_text_from_file(file, raw)

    share_id = None
    if visibility == "shared":
        share_id = uuid.uuid4().hex[:12]

    doc = DocumentSchema(
        owner_id=str(user.get("_id")),
        filename=file.filename,
        mime_type=file.content_type,
        size=len(raw),
        text=text,
        visibility="shared" if visibility == "shared" else "private",
        share_id=share_id,
    )
    doc_id = create_document("document", doc)
    return {"document_id": doc_id, "extracted_chars": len(text)}


# ------------------------
# Summarization via external models
# ------------------------
class SummarizeBody(BaseModel):
    document_id: str
    style: Optional[str] = "concise"
    custom_prompt: Optional[str] = None
    provider: Optional[str] = "openai-compatible"


def call_free_llm(prompt: str) -> str:
    # Fallback to a free endpoint. Here, we try OpenRouter's free models if key provided,
    # else return a simple heuristic summary. In production, set OPENROUTER_API_KEY or GEMINI_API_KEY.
    openrouter_key = os.getenv("OPENROUTER_API_KEY")
    if openrouter_key:
        try:
            resp = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {openrouter_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "google/gemma-2-9b-it:free",
                    "messages": [
                        {"role": "system", "content": "You are a precise legal-business summarizer. Return clean text only."},
                        {"role": "user", "content": prompt},
                    ],
                    "temperature": 0.2,
                },
                timeout=30,
            )
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        except Exception:
            pass
    # Heuristic fallback
    lines = [l.strip() for l in prompt.splitlines() if l.strip()]
    preview = " ".join(lines)[:800]
    return f"Summary (heuristic): {preview[:200]}..."


def style_instruction(style: str) -> str:
    mapping = {
        "concise": "Summarize concisely in 5-7 bullet points.",
        "detailed": "Provide a detailed summary with sections: Overview, Key Points, Risks, Next Steps.",
        "bullets": "Return only bullet points (max 12), each under 20 words.",
        "legal": "Summarize focusing on obligations, rights, liabilities, jurisdictions, and deadlines.",
        "executive": "Create an executive brief: context, impact, recommendations in 150-250 words.",
    }
    return mapping.get(style, mapping["concise"])


@app.post("/summaries/create")
def create_summary(body: SummarizeBody, token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    doc = db["document"].find_one({"_id": ObjectId(body.document_id), "owner_id": str(user.get("_id"))})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    style_note = style_instruction(body.style or "concise")
    prompt = f"Document text:\n\n{doc.get('text','')[:8000]}\n\nTask: {style_note}\nIf helpful, propose 3-5 follow-up questions."
    content = call_free_llm(prompt)

    sum_doc = SummarySchema(
        owner_id=str(user.get("_id")),
        document_id=str(doc.get("_id")),
        style=body.style or "concise",
        prompt=body.custom_prompt,
        content=content,
        model="openrouter:gemma-2-9b-it:free" if os.getenv("OPENROUTER_API_KEY") else "heuristic",
    )
    sum_id = create_document("summary", sum_doc)
    return {"summary_id": sum_id, "content": content}


@app.get("/summaries/by-document/{doc_id}")
def list_summaries(doc_id: str, token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    items = get_documents("summary", {"document_id": doc_id, "owner_id": str(user.get("_id"))}, limit=50)
    for it in items:
        it["_id"] = str(it["_id"])  # stringifying for JSON
    return items


# ------------------------
# Export endpoints
# ------------------------
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document as DocxWriter


@app.get("/export/txt/{summary_id}")
def export_txt(summary_id: str, token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    s = db["summary"].find_one({"_id": ObjectId(summary_id), "owner_id": str(user.get("_id"))})
    if not s:
        raise HTTPException(status_code=404, detail="Summary not found")
    return {"content": s.get("content", "")}


@app.get("/export/pdf/{summary_id}")
def export_pdf(summary_id: str, token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    s = db["summary"].find_one({"_id": ObjectId(summary_id), "owner_id": str(user.get("_id"))})
    if not s:
        raise HTTPException(status_code=404, detail="Summary not found")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    textobject = c.beginText(40, height - 50)
    textobject.setFont("Helvetica", 11)

    title = "Gamma Summary"
    for line in [title, "", f"Style: {s.get('style','concise')}", "", *s.get("content", "").splitlines()]:
        if textobject.getY() < 40:
            c.drawText(textobject)
            c.showPage()
            textobject = c.beginText(40, height - 50)
            textobject.setFont("Helvetica", 11)
        textobject.textLine(line)

    c.drawText(textobject)
    c.showPage()
    c.save()
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(buf, media_type="application/pdf", headers={
        "Content-Disposition": f"attachment; filename=summary_{summary_id}.pdf"
    })


@app.get("/export/docx/{summary_id}")
def export_docx(summary_id: str, token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    s = db["summary"].find_one({"_id": ObjectId(summary_id), "owner_id": str(user.get("_id"))})
    if not s:
        raise HTTPException(status_code=404, detail="Summary not found")

    doc = DocxWriter()
    doc.add_heading("Gamma Summary", level=1)
    doc.add_paragraph(f"Style: {s.get('style','concise')}")
    doc.add_paragraph("")
    for line in s.get("content", "").splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
        "Content-Disposition": f"attachment; filename=summary_{summary_id}.docx"
    })


# ------------------------
# Simple analytics & privacy toggles
# ------------------------
@app.get("/analytics/basic")
def basic_analytics(token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    uid = str(user.get("_id"))
    docs = db["document"].count_documents({"owner_id": uid})
    sums = db["summary"].count_documents({"owner_id": uid})
    chars = sum((d.get("size", 0) for d in db["document"].find({"owner_id": uid})) or [0])
    return {"documents": docs, "summaries": sums, "uploaded_bytes": chars}


class PrivacyBody(BaseModel):
    visibility: str


@app.post("/documents/{doc_id}/privacy")
def set_privacy(doc_id: str, body: PrivacyBody, token: str = Depends(oauth2_scheme)):
    user = get_current_user(token)
    vis = body.visibility if body.visibility in ("private", "shared") else "private"
    share_id = uuid.uuid4().hex[:12] if vis == "shared" else None
    res = db["document"].find_one_and_update(
        {"_id": ObjectId(doc_id), "owner_id": str(user.get("_id"))},
        {"$set": {"visibility": vis, "share_id": share_id, "updated_at": datetime.now(timezone.utc)}},
        return_document=ReturnDocument.AFTER,
    )
    if not res:
        raise HTTPException(status_code=404, detail="Document not found")
    res["_id"] = str(res["_id"])  # jsonify id
    return res


# Public read for shared documents (no auth)
@app.get("/shared/{share_id}")
def get_shared_document(share_id: str):
    doc = db["document"].find_one({"share_id": share_id, "visibility": "shared"})
    if not doc:
        raise HTTPException(status_code=404, detail="Not found")
    doc["_id"] = str(doc["_id"])  # jsonify id
    # Do not expose raw text for privacy; only meta
    return {"filename": doc.get("filename"), "mime_type": doc.get("mime_type"), "size": doc.get("size")}


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
